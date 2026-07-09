from io import BytesIO

import docx
from resume_exporter import export_docx
from resume_template import ResumeTemplate

RESUME_TEXT = """John Doe
john.doe@email.com | 555-123-4567

WORK EXPERIENCE
Software Engineer, Acme Corp (2021-2024)
Remote
Built ETL pipelines using Python and SQL.

EDUCATION
BS Computer Science, State University (2017-2021)

SKILLS
Languages: Python, SQL

SUMMARY
Backend engineer with data pipeline experience.

CERTIFICATIONS
AWS Certified Solutions Architect
"""


def _paragraph_runs(paragraph):
    return [(run.text, bool(run.bold), bool(run.italic)) for run in paragraph.runs]


def test_produces_a_valid_reopenable_docx():
    data = export_docx(RESUME_TEXT)
    assert isinstance(data, bytes) and len(data) > 0
    doc = docx.Document(BytesIO(data))
    assert len(doc.paragraphs) > 0


def test_name_is_bold_and_centered():
    doc = docx.Document(BytesIO(export_docx(RESUME_TEXT)))
    name_paragraph = doc.paragraphs[0]
    assert name_paragraph.text == "John Doe"
    assert name_paragraph.runs[0].bold is True


def test_experience_entry_has_bold_heading_with_tab_separated_dates():
    doc = docx.Document(BytesIO(export_docx(RESUME_TEXT)))
    heading_paragraph = next(p for p in doc.paragraphs if "Acme Corp" in p.text)
    assert heading_paragraph.text == "Software Engineer, Acme Corp\t2021-2024"
    runs = _paragraph_runs(heading_paragraph)
    assert runs[0] == ("Software Engineer, Acme Corp", True, False)
    assert runs[1] == ("\t2021-2024", False, False)


def test_subheading_is_italic():
    doc = docx.Document(BytesIO(export_docx(RESUME_TEXT)))
    subheading_paragraph = next(p for p in doc.paragraphs if p.text == "Remote")
    assert subheading_paragraph.runs[0].italic is True


def test_bullets_use_list_bullet_style():
    doc = docx.Document(BytesIO(export_docx(RESUME_TEXT)))
    bullet_paragraph = next(p for p in doc.paragraphs if "Built ETL pipelines" in p.text)
    assert bullet_paragraph.style.name == "List Bullet"


def test_skills_label_is_bold_rest_is_not():
    doc = docx.Document(BytesIO(export_docx(RESUME_TEXT)))
    skills_paragraph = next(p for p in doc.paragraphs if p.text.startswith("Languages"))
    runs = _paragraph_runs(skills_paragraph)
    assert runs[0] == ("Languages: ", True, False)
    assert runs[1] == ("Python, SQL", False, False)


def test_unmapped_section_appended_with_header_and_bullets():
    doc = docx.Document(BytesIO(export_docx(RESUME_TEXT)))
    paragraph_texts = [p.text for p in doc.paragraphs]
    certifications_index = paragraph_texts.index("CERTIFICATIONS")
    assert paragraph_texts[certifications_index + 1] == "AWS Certified Solutions Architect"


def test_summary_renders_as_single_paragraph_before_experience():
    doc = docx.Document(BytesIO(export_docx(RESUME_TEXT)))
    paragraph_texts = [p.text for p in doc.paragraphs]
    summary_index = paragraph_texts.index("SUMMARY")
    experience_index = paragraph_texts.index("EXPERIENCE")
    assert summary_index < experience_index

    summary_paragraph = doc.paragraphs[summary_index + 1]
    assert summary_paragraph.text == "Backend engineer with data pipeline experience."
    assert summary_paragraph.style.name != "List Bullet"


def test_missing_section_is_not_rendered():
    doc = docx.Document(BytesIO(export_docx(RESUME_TEXT)))
    paragraph_texts = [p.text for p in doc.paragraphs]
    assert "PROJECTS" not in paragraph_texts


def test_falls_back_to_plain_paragraphs_when_no_sections_found():
    text = "Just plain text\nwith no section headers\n"
    doc = docx.Document(BytesIO(export_docx(text)))
    paragraph_texts = [p.text for p in doc.paragraphs]
    assert paragraph_texts == ["Just plain text", "with no section headers"]


def test_empty_input_produces_empty_document():
    doc = docx.Document(BytesIO(export_docx("")))
    assert len(doc.paragraphs) == 0


def test_dateless_project_title_renders_bold():
    text = (
        "John Doe\n"
        "john.doe@email.com\n"
        "\n"
        "PROJECTS\n"
        "ResuMatch | Python, Streamlit\n"
        "Built an LLM-powered resume-tailoring app that aligns resumes to job descriptions.\n"
    )
    doc = docx.Document(BytesIO(export_docx(text)))
    heading_paragraph = next(p for p in doc.paragraphs if p.text == "ResuMatch | Python, Streamlit")
    assert heading_paragraph.runs[0].bold is True


def test_custom_template_changes_section_order_and_name_style():
    custom = ResumeTemplate(
        section_order=(
            ("SKILLS", "skills"),
            ("EXPERIENCE", "experience"),
            ("PROJECTS", "projects"),
            ("EDUCATION", "education"),
        ),
        name_bold=False,
    )
    doc = docx.Document(BytesIO(export_docx(RESUME_TEXT, template=custom)))
    paragraph_texts = [p.text for p in doc.paragraphs]

    assert paragraph_texts.index("SKILLS") < paragraph_texts.index("EXPERIENCE")
    assert doc.paragraphs[0].runs[0].bold is False
