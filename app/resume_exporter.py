from io import BytesIO

import docx
from docx.enum.text import WD_TAB_ALIGNMENT
from resume_structurer import structure_resume
from resume_template import DEFAULT_TEMPLATE, ResumeTemplate


def _add_section_header(document, title, template):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = template.section_header_space_before
    paragraph.paragraph_format.space_after = template.section_header_space_after
    run = paragraph.add_run(title)
    run.bold = template.section_header_bold
    run.font.size = template.section_header_font_size
    run.font.small_caps = template.section_header_small_caps


def _add_heading_dates_paragraph(document, heading, dates, template):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        template.date_tab_position, alignment=WD_TAB_ALIGNMENT.RIGHT
    )
    if heading:
        paragraph.add_run(heading).bold = template.entry_heading_bold
    if dates:
        paragraph.add_run("\t" + dates)


def _add_subheading_paragraph(document, subheading, template):
    paragraph = document.add_paragraph()
    paragraph.add_run(subheading).italic = template.subheading_italic


def _add_bullets(document, bullets, template):
    for bullet in bullets:
        paragraph = document.add_paragraph(bullet, style=template.bullet_style)
        if paragraph.runs:
            paragraph.runs[0].font.size = template.body_font_size


def _add_summary_paragraph(document, lines, template):
    paragraph = document.add_paragraph(" ".join(lines))
    paragraph.runs[0].font.size = template.body_font_size


def _add_entries(document, entries, template):
    for entry in entries:
        if entry["heading"] or entry["dates"]:
            _add_heading_dates_paragraph(document, entry["heading"], entry["dates"], template)
        if entry["subheading"]:
            _add_subheading_paragraph(document, entry["subheading"], template)
        _add_bullets(document, entry["bullets"], template)


def _add_skill_line(document, line, template):
    paragraph = document.add_paragraph()
    if ":" in line:
        label, _, rest = line.partition(":")
        paragraph.add_run(label.strip() + ": ").bold = template.skills_label_bold
        paragraph.add_run(rest.strip())
    else:
        paragraph.add_run(line)


def _has_recognized_sections(structured: dict) -> bool:
    """Whether structure_resume found at least one real section. Deliberately ignores name/
    contact: those are the only two things structure_resume can extract with no headers at
    all, and rendering just a name with nothing else would silently drop the rest of the
    resume -- better to fall back to a plain per-line dump when no real structure was found."""
    return bool(
        structured["summary"]
        or structured["experience"]
        or structured["projects"]
        or structured["education"]
        or structured["skills"]
        or structured["other_sections"]
    )


def export_docx(tailored_text: str, template: ResumeTemplate = DEFAULT_TEMPLATE) -> bytes:
    """
    Builds a fixed-template DOCX from tailored_text: a centered bold name and contact line,
    then SUMMARY / EXPERIENCE / PROJECTS / EDUCATION / SKILLS sections (only those with
    content), with each experience/project/education entry rendered as a bold heading with a
    right-aligned date on the same line, an italic subheading line, and bulleted achievements.
    SUMMARY renders as a single flowing paragraph rather than bullets. Any other recognized
    section from the original resume (e.g. Certifications) that doesn't map to those five is
    appended at the end, so nothing from the original resume is silently dropped. All fonts,
    sizes, spacing, and section order come from `template` (see resume_template.py) — swap in a
    different ResumeTemplate to change the look without touching this rendering logic.
    Structure is extracted heuristically (regex date-range detection + line position) via
    resume_structurer.structure_resume — resumes formatted very differently than expected may
    not parse into clean entries, in which case this falls back to one plain paragraph per line.
    """
    structured = structure_resume(tailored_text)
    document = docx.Document()
    document.styles["Normal"].font.name = template.font_name
    section = document.sections[0]
    section.left_margin = template.page_margin
    section.right_margin = template.page_margin
    section.top_margin = template.page_margin
    section.bottom_margin = template.page_margin

    if not _has_recognized_sections(structured):
        for line in tailored_text.splitlines():
            document.add_paragraph(line)
    else:
        if structured["name"]:
            name_paragraph = document.add_paragraph()
            name_paragraph.alignment = template.name_alignment
            run = name_paragraph.add_run(structured["name"])
            run.bold = template.name_bold
            run.font.size = template.name_font_size
            run.font.small_caps = template.name_small_caps

        if structured["contact"]:
            contact_paragraph = document.add_paragraph()
            contact_paragraph.alignment = template.contact_alignment
            contact_run = contact_paragraph.add_run(structured["contact"])
            contact_run.font.size = template.contact_font_size

        for title, key in template.section_order:
            if key == "skills":
                if structured["skills"]:
                    _add_section_header(document, title, template)
                    for line in structured["skills"]:
                        _add_skill_line(document, line, template)
            elif key == "summary":
                if structured["summary"]:
                    _add_section_header(document, title, template)
                    _add_summary_paragraph(document, structured["summary"], template)
            elif structured[key]:
                _add_section_header(document, title, template)
                _add_entries(document, structured[key], template)

        for header_title, lines in structured["other_sections"]:
            _add_section_header(document, header_title, template)
            _add_bullets(document, lines, template)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
